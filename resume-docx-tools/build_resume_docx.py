"""Build a clean, ATS-safe .docx resume from a Markdown source file.

Why: Microsoft's stock resume templates put the entire body inside a layout
table and wrap the name and work history in content controls (w:sdt). Standard
parsers - including the extraction layer of many applicant-tracking systems -
can silently drop exactly those parts: a plain python-docx read of such a file
returns no candidate name and an empty work history. This builder produces the
opposite: plain paragraphs, real Heading 1/2/3 styles, genuine Word list
bullets, working hyperlinks, no tables, no content controls, clean metadata.

Markdown dialect (one construct per line):

    # Name                              -> large name line (first H1 only)
    **Headline | keywords**             -> bold headline (first bold-only line)
    plain text with [links](url)        -> contact line / paragraph
    ## Section                          -> Heading 1
    ### Role | Employer - Location      -> Heading 2 (job entries)
    #### Project name (dates)           -> Heading 3 (projects within a role)
    *March 2026 - Present · note*       -> small italic date line
    - **Label:** text                   -> bullet with bold label
    - text                              -> bullet

Inline **bold** and [text](url) hyperlinks work everywhere.

Usage:
    python build_resume_docx.py INPUT.md [-o OUTPUT.docx]
                                [--title TEXT] [--author NAME] [--check]

--check verifies the output: every visible Markdown line must come back from a
plain-text extraction of the .docx (the exact test stock templates fail).

Requires: python-docx  (pip install python-docx)
"""
from __future__ import annotations

import argparse
import html
import re
import sys
import zipfile
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x21, 0x21, 0x21)
ACCENT = RGBColor(0x1F, 0x3B, 0x66)

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def add_hyperlink(paragraph, text, url, size=None, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    r = paragraph._p.makeelement(qn("w:r"), {})
    rPr = paragraph._p.makeelement(qn("w:rPr"), {})
    rPr.append(paragraph._p.makeelement(qn("w:color"), {qn("w:val"): "1F3B66"}))
    rPr.append(paragraph._p.makeelement(qn("w:u"), {qn("w:val"): "single"}))
    if size:
        rPr.append(paragraph._p.makeelement(qn("w:sz"), {qn("w:val"): str(int(size * 2))}))
    if bold:
        rPr.append(paragraph._p.makeelement(qn("w:b"), {}))
    r.append(rPr)
    t = paragraph._p.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def add_rich_text(paragraph, text, size=None, base_bold=False):
    pos = 0
    for m in LINK_RE.finditer(text):
        _emit_bold_aware(paragraph, text[pos:m.start()], size, base_bold)
        add_hyperlink(paragraph, m.group(1), m.group(2), size=size, bold=base_bold)
        pos = m.end()
    _emit_bold_aware(paragraph, text[pos:], size, base_bold)


def _emit_bold_aware(paragraph, text, size, base_bold):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], size, base_bold)
        _run(paragraph, m.group(1), size, True)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], size, base_bold)


def _run(paragraph, text, size, bold):
    if not text:
        return
    r = paragraph.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)


def style_setup(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.space_before = Pt(0)
    for name, size, color, before, after in [
        ("Heading 1", 13, ACCENT, 10, 4),
        ("Heading 2", 11.5, INK, 8, 2),
        ("Heading 3", 10.5, INK, 6, 2),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(10.5)
    lb.font.color.rgb = INK
    lb.paragraph_format.space_after = Pt(3)
    lb.paragraph_format.left_indent = Inches(0.25)


def build(md_path: Path, docx_path: Path, title: str | None, author: str | None):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = docx.Document()
    style_setup(doc)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.6)
        section.left_margin = section.right_margin = Inches(0.75)

    name_text = None
    seen_headline = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if line.startswith("# ") and name_text is None:
            name_text = line[2:].strip()
            p = doc.add_paragraph()
            _run(p, name_text, 22, True)
            p.paragraph_format.space_after = Pt(1)
        elif line.startswith("#### "):
            add_rich_text(doc.add_paragraph(style="Heading 3"), line[5:].strip(), base_bold=True)
        elif line.startswith("### "):
            add_rich_text(doc.add_paragraph(style="Heading 2"), line[4:].strip(), base_bold=True)
        elif line.startswith("## "):
            add_rich_text(doc.add_paragraph(style="Heading 1"), line[3:].strip(), base_bold=True)
        elif line.startswith("- "):
            add_rich_text(doc.add_paragraph(style="List Bullet"), line[2:].strip())
        elif re.fullmatch(r"\*[^*].*\*", stripped):
            p = doc.add_paragraph()
            r = p.add_run(stripped[1:-1])
            r.italic = True
            r.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(4)
        elif not seen_headline and BOLD_RE.fullmatch(stripped):
            p = doc.add_paragraph()
            add_rich_text(p, stripped[2:-2], size=11.5, base_bold=True)
            p.paragraph_format.space_after = Pt(2)
            seen_headline = True
        else:
            add_rich_text(doc.add_paragraph(), stripped)

    cp = doc.core_properties
    cp.title = title or (name_text or docx_path.stem)
    if author or name_text:
        cp.author = author or name_text
        cp.last_modified_by = author or name_text
    for attr in ("category", "comments", "keywords", "subject"):
        setattr(cp, attr, "")
    doc.save(str(docx_path))
    return name_text


def md_visible_lines(md_path: Path) -> list[str]:
    out = []
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"^#{1,6}\s+", "", line)
        line = re.sub(r"^- ", "", line)
        if re.fullmatch(r"\*[^*].*\*", line):
            line = line[1:-1]
        line = LINK_RE.sub(lambda m: m.group(1), line)
        line = line.replace("**", "")
        out.append(re.sub(r"\s+", " ", line).strip())
    return out


def docx_visible_lines(docx_path: Path) -> list[str]:
    xml = zipfile.ZipFile(docx_path).read("word/document.xml").decode("utf-8")
    out = []
    for p in re.split(r"</w:p>", xml):
        text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p))
        text = re.sub(r"\s+", " ", html.unescape(text)).strip()
        if text:
            out.append(text)
    return out


def check(md_path: Path, docx_path: Path, name_text: str | None) -> list[str]:
    problems = []
    extracted = "\n".join(p.text for p in docx.Document(str(docx_path)).paragraphs)
    if name_text and name_text not in extracted:
        problems.append(f"parser test FAILED: name {name_text!r} not extracted")
    md_lines = md_visible_lines(md_path)
    dx_lines = docx_visible_lines(docx_path)
    if md_lines != dx_lines:
        for x in [l for l in md_lines if l not in dx_lines][:5]:
            problems.append(f"parity: markdown-only line: {x[:90]}")
        for x in [l for l in dx_lines if l not in md_lines][:5]:
            problems.append(f"parity: docx-only line: {x[:90]}")
    return problems


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("input", help="Markdown resume source")
    ap.add_argument("-o", "--output", help="output .docx (default: input name with .docx)")
    ap.add_argument("--title", help="document title metadata (default: the name from the first H1)")
    ap.add_argument("--author", help="author metadata (default: the name from the first H1)")
    ap.add_argument("--check", action="store_true", help="verify parser extraction and md/docx parity")
    args = ap.parse_args()

    md_path = Path(args.input)
    docx_path = Path(args.output) if args.output else md_path.with_suffix(".docx")
    name_text = build(md_path, docx_path, args.title, args.author)
    words = sum(len(l.split()) for l in docx_visible_lines(docx_path))
    print(f"built {docx_path} (~{words} words)")
    if args.check:
        problems = check(md_path, docx_path, name_text)
        for p in problems:
            print("   " + p)
        print("   check: " + ("clean" if not problems else f"{len(problems)} problem(s)"))
        return 0 if not problems else 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
