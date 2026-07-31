"""Document ingestion for the local RAG index: text extraction and chunking.

Extraction accepts UTF-8 text files, .pdf (text layer, via pypdf), and .docx
(via python-docx). Chunking is character-based with boundary-aware breaks so
no tokenizer dependency is needed. Both steps are pure and offline; embedding
happens in server.py.
"""

from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from mcp_chatbot import attachments

SUPPORTED_FORMATS = "UTF-8 text files, .pdf (with a text layer), and .docx"
MAX_DOCUMENT_BYTES = attachments.MAX_ATTACHMENT_BYTES

# ~400 English tokens per chunk: precise retrieval, far under embedding input
# limits even for dense scripts. Overlap preserves context across boundaries.
CHUNK_CHARS = 1600
CHUNK_OVERLAP = 200

_SENTENCE_END = re.compile(r"(?<=[.!?])\s")


def extract_text(path_str: str) -> tuple[Path, str, str]:
    """Extract plain text from a document for indexing.

    Returns (resolved path, extracted text, sha256 of the raw file bytes). The
    hash covers the bytes rather than the extraction so unchanged files can be
    skipped on re-upload without depending on parser stability.
    """
    path = Path(path_str)
    if not path.is_file():
        raise ValueError(f"Document not found: {path_str}")
    try:
        size = path.stat().st_size
        path = path.resolve()
        raw = path.read_bytes()
    except OSError as exc:
        # PermissionError for files locked by another application, or the file
        # vanished after the is_file() check; either way a per-file problem.
        raise ValueError(
            f"Document {path_str} could not be read ({exc}); check that it still "
            "exists and is not locked by another application, then retry."
        ) from exc
    if size > MAX_DOCUMENT_BYTES:
        raise ValueError(
            f"Document too large (over {MAX_DOCUMENT_BYTES // (1024 * 1024)} MB): {path_str}"
        )
    digest = hashlib.sha256(raw).hexdigest()
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _pdf_text(path, raw)
    elif suffix == ".docx":
        text = _docx_text(path, raw)
    elif suffix == ".doc":
        raise ValueError(
            f"Legacy .doc is not supported: {path}. Save it as .docx first "
            f"(supported formats: {SUPPORTED_FORMATS})."
        )
    else:
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError(
                f"Document {path} is not UTF-8 text; supported formats: {SUPPORTED_FORMATS}."
            ) from exc
    if not text.strip():
        raise ValueError(
            f"No extractable text in {path.name} - a scanned PDF without a text "
            "layer cannot be indexed; run OCR first."
        )
    return path, text, digest


def _pdf_text(path: Path, raw: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(raw))
        # Some PDFs are "encrypted" with an empty owner password and still open;
        # anything needing a real password cannot be indexed.
        locked = reader.is_encrypted and not reader.decrypt("")
    except Exception as exc:  # pypdf raises a zoo of parse errors
        raise ValueError(f"Could not extract text from PDF {path}: {exc}") from exc
    if locked:
        raise ValueError(f"Document {path} is password-protected; decrypt it first.")
    try:
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise ValueError(f"Could not extract text from PDF {path}: {exc}") from exc


def _docx_text(path: Path, raw: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(raw))
    except Exception as exc:  # bad zip, not a docx, truncated file
        raise ValueError(f"Could not extract text from docx {path}: {exc}") from exc
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n\n".join(part for part in parts if part.strip())


def chunk_text(
    text: str, size: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP
) -> list[dict[str, Any]]:
    """Split text into overlapping chunks, preferring natural boundaries.

    Returns [{"text", "index", "start", "end"}, ...] where start/end are
    offsets into the CRLF-normalized text, so ``normalized[start:end] ==
    chunk["text"]`` always holds. Whitespace-only windows are skipped and
    indices stay sequential.
    """
    if size <= overlap:
        raise ValueError(f"Chunk size ({size}) must exceed overlap ({overlap}).")
    text = text.replace("\r\n", "\n")
    chunks: list[dict[str, Any]] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + size, length)
        if end < length:
            end = _best_break(text, start, end)
        piece = text[start:end]
        stripped = piece.strip()
        if stripped:
            lead = len(piece) - len(piece.lstrip())
            chunks.append(
                {
                    "text": stripped,
                    "index": len(chunks),
                    "start": start + lead,
                    "end": start + lead + len(stripped),
                }
            )
        if end >= length:
            break
        # max() guarantees forward progress even on boundary-free text.
        start = max(end - overlap, start + 1)
    return chunks


def _best_break(text: str, start: int, end: int) -> int:
    """Best cut position for the window [start, end): paragraph > newline >
    sentence end > word break > hard cut.

    Only the back half of the window is searched, so boundary placement can
    never shrink a chunk below half its nominal size.
    """
    floor = start + (end - start) // 2
    for sep in ("\n\n", "\n"):
        pos = text.rfind(sep, floor, end)
        if pos != -1:
            return pos
    matches = list(_SENTENCE_END.finditer(text, floor, end))
    if matches:
        return matches[-1].start()
    pos = text.rfind(" ", floor, end)
    if pos != -1:
        return pos
    return end
