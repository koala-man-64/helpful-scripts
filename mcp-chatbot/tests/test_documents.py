"""Extraction and chunking tests for mcp_chatbot.documents (pure, offline)."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

import docx
import pytest

from mcp_chatbot import documents

# --- extract_text ---


def test_utf8_text_extracted_with_hash(tmp_path: Path) -> None:
    f = tmp_path / "notes.txt"
    f.write_text("hello world", encoding="utf-8")
    path, text, digest = documents.extract_text(str(f))
    assert path == f.resolve()
    assert text == "hello world"
    assert len(digest) == 64  # sha256 hex


def test_missing_document_rejected() -> None:
    with pytest.raises(ValueError, match="Document not found"):
        documents.extract_text("no-such-file.txt")


def test_oversized_document_rejected(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(documents, "MAX_DOCUMENT_BYTES", 4)
    f = tmp_path / "big.txt"
    f.write_text("hello world", encoding="utf-8")
    with pytest.raises(ValueError, match="too large"):
        documents.extract_text(str(f))


def test_unreadable_document_rejected(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    f = tmp_path / "locked.txt"
    f.write_text("data", encoding="utf-8")

    def boom(self: Path) -> bytes:
        raise PermissionError("locked by another process")

    monkeypatch.setattr(Path, "read_bytes", boom)
    with pytest.raises(ValueError, match="could not be read"):
        documents.extract_text(str(f))


def test_non_utf8_rejected(tmp_path: Path) -> None:
    f = tmp_path / "blob.dat"
    f.write_bytes(b"\xff\xfe\x01\x02")
    with pytest.raises(ValueError, match="not UTF-8"):
        documents.extract_text(str(f))


def test_whitespace_only_text_rejected(tmp_path: Path) -> None:
    f = tmp_path / "empty.txt"
    f.write_text("   \n\n  ", encoding="utf-8")
    with pytest.raises(ValueError, match="No extractable text"):
        documents.extract_text(str(f))


def test_legacy_doc_rejected(tmp_path: Path) -> None:
    f = tmp_path / "old.doc"
    f.write_bytes(b"junk")
    with pytest.raises(ValueError, match=r"Legacy \.doc"):
        documents.extract_text(str(f))


def test_docx_paragraphs_and_tables_extracted(tmp_path: Path) -> None:
    d = docx.Document()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    table = d.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "alpha"
    table.cell(0, 1).text = "beta"
    f = tmp_path / "memo.docx"
    d.save(str(f))
    _, text, _ = documents.extract_text(str(f))
    assert "First paragraph." in text
    assert "Second paragraph." in text
    assert "alpha\tbeta" in text


def test_corrupt_docx_rejected(tmp_path: Path) -> None:
    f = tmp_path / "bad.docx"
    f.write_bytes(b"this is not a zip archive")
    with pytest.raises(ValueError, match="Could not extract text from docx"):
        documents.extract_text(str(f))


def _fake_pdf(monkeypatch: pytest.MonkeyPatch, pages: list[str], encrypted: bool = False) -> None:
    reader = SimpleNamespace(
        is_encrypted=encrypted,
        decrypt=lambda password: 0,  # 0 = pypdf PasswordType.NOT_DECRYPTED
        pages=[SimpleNamespace(extract_text=lambda t=t: t) for t in pages],
    )
    monkeypatch.setattr(documents, "PdfReader", lambda buffer: reader)


def test_pdf_pages_joined(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _fake_pdf(monkeypatch, ["Page one.", "Page two."])
    f = tmp_path / "doc.pdf"
    f.write_bytes(b"%PDF-1.4 fake")
    _, text, _ = documents.extract_text(str(f))
    assert text == "Page one.\n\nPage two."


def test_password_protected_pdf_rejected(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _fake_pdf(monkeypatch, ["secret"], encrypted=True)
    f = tmp_path / "locked.pdf"
    f.write_bytes(b"%PDF-1.4 fake")
    with pytest.raises(ValueError, match="password-protected"):
        documents.extract_text(str(f))


def test_pdf_without_text_layer_rejected(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _fake_pdf(monkeypatch, ["", ""])  # scanned pages: no text layer
    f = tmp_path / "scan.pdf"
    f.write_bytes(b"%PDF-1.4 fake")
    with pytest.raises(ValueError, match="No extractable text"):
        documents.extract_text(str(f))


def test_corrupt_pdf_rejected_by_real_parser(tmp_path: Path) -> None:
    f = tmp_path / "junk.pdf"
    f.write_bytes(b"not a pdf at all")
    with pytest.raises(ValueError, match="Could not extract text from PDF"):
        documents.extract_text(str(f))


# --- chunk_text ---


def test_short_text_is_single_chunk() -> None:
    assert documents.chunk_text("hello world") == [
        {"text": "hello world", "index": 0, "start": 0, "end": 11}
    ]


def test_offsets_reproduce_chunk_text_and_windows_overlap() -> None:
    text = "word " * 1000  # 5000 chars, word breaks only
    chunks = documents.chunk_text(text, size=1000, overlap=100)
    assert len(chunks) > 1
    for chunk in chunks:
        assert text[chunk["start"] : chunk["end"]] == chunk["text"]
    assert [c["index"] for c in chunks] == list(range(len(chunks)))
    for prev, nxt in zip(chunks, chunks[1:]):
        assert nxt["start"] < prev["end"]  # consecutive windows overlap
    assert chunks[-1]["end"] == len(text.rstrip())


def test_paragraph_boundary_preferred() -> None:
    text = "a" * 900 + "\n\n" + "b" * 900
    chunks = documents.chunk_text(text, size=1000, overlap=100)
    assert chunks[0]["text"] == "a" * 900


def test_sentence_boundary_preferred() -> None:
    text = "x" * 700 + ". " + "y" * 700
    chunks = documents.chunk_text(text, size=1000, overlap=100)
    assert chunks[0]["text"] == "x" * 700 + "."


def test_boundary_free_text_still_progresses() -> None:
    text = "z" * 3500
    chunks = documents.chunk_text(text, size=1000, overlap=100)
    assert chunks[0] == {"text": "z" * 1000, "index": 0, "start": 0, "end": 1000}
    assert chunks[-1]["end"] == 3500


def test_crlf_offsets_refer_to_normalized_text() -> None:
    assert documents.chunk_text("line one\r\nline two") == [
        {"text": "line one\nline two", "index": 0, "start": 0, "end": 17}
    ]


def test_whitespace_only_windows_skipped() -> None:
    text = "a" * 10 + " " * 2000 + "b" * 10
    chunks = documents.chunk_text(text, size=1000, overlap=100)
    assert [c["text"] for c in chunks] == ["a" * 10, "b" * 10]
    assert [c["index"] for c in chunks] == [0, 1]


def test_size_must_exceed_overlap() -> None:
    with pytest.raises(ValueError, match="must exceed overlap"):
        documents.chunk_text("abc", size=100, overlap=100)
